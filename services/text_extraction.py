"""Extracts plain text from uploaded resume files (PDF / DOCX)."""
import io
from pathlib import Path

import docx
from pypdf import PdfReader

try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:  # pragma: no cover - defensive fallback
    _HAS_PDFPLUMBER = False

from services.exceptions import TextExtractionError, UnsupportedFileTypeError


def _extract_with_pdfplumber(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        pages_text = [page.extract_text() or "" for page in pdf.pages]
    return "\n".join(pages_text).strip()


def _extract_with_pypdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages_text).strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    # pdfplumber generally preserves line breaks and layout (columns, bullet
    # lists) more faithfully than pypdf, which matters for the section- and
    # date-line-based parsing this app relies on. Try it first, and fall
    # back to pypdf for PDFs pdfplumber can't handle.
    text = ""
    errors = []

    if _HAS_PDFPLUMBER:
        try:
            text = _extract_with_pdfplumber(file_bytes)
        except Exception as exc:  # noqa: BLE001
            errors.append(str(exc))

    if not text:
        try:
            text = _extract_with_pypdf(file_bytes)
        except Exception as exc:  # noqa: BLE001
            errors.append(str(exc))

    if not text:
        raise TextExtractionError(
            "No extractable text found in the PDF.",
            details="The PDF may be a scanned image without a text layer."
            + (f" Errors: {'; '.join(errors)}" if errors else ""),
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        text = "\n".join(p for p in paragraphs if p and p.strip()).strip()
        if not text:
            raise TextExtractionError("No extractable text found in the DOCX file.")
        return text
    except TextExtractionError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise TextExtractionError("Failed to read the DOCX file.", details=str(exc)) from exc


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if suffix == ".docx":
        return extract_text_from_docx(file_bytes)
    raise UnsupportedFileTypeError(
        f"Unsupported file type '{suffix}'. Only .pdf and .docx are supported."
    )
